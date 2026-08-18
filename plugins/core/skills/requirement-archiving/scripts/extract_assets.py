# -*- coding: utf-8 -*-
"""历史需求收料提取：模态盘点 + 正文/图片/外链全量无损提取。

用法:
    python extract_assets.py <源目录> <工作目录>

产出（写入 <工作目录>）:
    inventory.md   模态矩阵：文件 × 类型 × 体量 × 正文字符数 × 图片数 × 外链数 × 备注
    text/*.txt     逐文件正文提取（docx 段落+表格；pdf 逐页；xls/xlsx 全表）
    images/*       内嵌图片无损提取（docx word/media 原件；PDF 内嵌位图 render=False 原始提取）
    links.md       外链清单（docx 超链接 rels + 各正文中的 http(s) 链接，标注出处）

依赖: python-docx, pdfplumber, pypdfium2, openpyxl, xlrd, Pillow
    一次安装: pip install python-docx pdfplumber pypdfium2 openpyxl xlrd Pillow
    缺库不中断收料——对应格式文件在 inventory 备注标明缺哪个库怎么装，装完重跑即可。
实战校准: 2026-07 两次大规模归档（16 份现状基线 PRD / 21 包 39 需求）。关键教训已固化——
    * PDF 内嵌图必须 render=False 原始位图提取，渲染放大是插值假高清；
    * xls 需 ignore_workbook_corruption 容错（TAPD 导出常见）；
    * 外链（尤其在线文档链接）是删源前必须清点的知情项。
"""
import sys
import io
import math
import os
import re
import shutil
import traceback
import warnings

warnings.filterwarnings("ignore")
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")

URL_RE = re.compile(r"https?://[^\s\)】」）,;；]+")

# 模块名 → pip 包名（缺库提示用；模块名与包名不一致的只有 docx 和 PIL）
PIP_NAMES = {"docx": "python-docx", "PIL": "Pillow"}


def safe_name(s, limit=80):
    return re.sub(r'[\\/:*?"<>|]', "_", s)[:limit]


def extract_docx(path, img_dir, tag):
    """返回 (正文, 图片数, 超链接列表, 图片写失败数)。图片按原件字节流拷出（无损）。"""
    import docx
    d = docx.Document(path)
    parts = []
    for p in d.paragraphs:
        t = p.text.strip()
        if t:
            style = p.style.name if p.style is not None else ""
            prefix = "[%s] " % style if style.lower().startswith("heading") else ""
            parts.append(prefix + t)
    for i, tb in enumerate(d.tables):
        parts.append("\n[TABLE %d]" % (i + 1))
        for row in tb.rows:
            parts.append(" | ".join(c.text.strip().replace("\n", " / ") for c in row.cells))
    n_img = 0
    img_fail = 0
    import zipfile
    z = zipfile.ZipFile(path)
    for info in z.infolist():
        if info.filename.startswith("word/media/") and info.file_size > 0:
            out = os.path.join(img_dir, "%s_%s" % (tag, os.path.basename(info.filename)))
            # 单张图写失败不能拖垮整份文件的正文：tag 现在带目录前缀，深目录 + 长文件名
            # 在 Windows 上可能撞 MAX_PATH；异常冒泡出去会让**已经提取好的正文**一起作废
            # （main 的 except 只留一句 note）。降级为计数，由调用方写进 inventory 备注。
            try:
                with open(out, "wb") as f:
                    f.write(z.read(info.filename))
                n_img += 1
            except OSError:
                img_fail += 1
    links = []
    for r in d.part.rels.values():
        if "hyperlink" in r.reltype and not r.target_ref.startswith("mailto"):
            links.append(r.target_ref)
    return "\n".join(parts), n_img, links


def extract_pdf(path, img_dir, tag):
    """返回 (正文, 图片数)。图片用 pypdfium2 render=False 提取原始位图。"""
    import pdfplumber
    parts = []
    with pdfplumber.open(path) as pdf:
        for pg in pdf.pages:
            parts.append(pg.extract_text() or "")
    n_img = 0
    try:
        import pypdfium2 as pdfium
        import pypdfium2.raw as pdfium_c
        doc = pdfium.PdfDocument(path)
        for pi in range(len(doc)):
            k = 0
            for obj in doc[pi].get_objects(filter=(pdfium_c.FPDF_PAGEOBJ_IMAGE,)):
                try:
                    pil = obj.get_bitmap(render=False).to_pil()
                except Exception:
                    continue
                if pil.width * pil.height < 40000:      # <200x200 视为装饰性小图
                    continue
                k += 1
                n_img += 1
                pil.save(os.path.join(img_dir, "%s_p%d_%d.png" % (tag, pi + 1, k)))
        doc.close()
    except ImportError:
        parts.append("\n[!! pypdfium2 不可用，内嵌图片未提取]")
    return "\n".join(parts), n_img


def extract_xls(path):
    import xlrd     # ImportError 须冒泡到 main 的缺库分支，不能被下面的兜底吞掉
    try:
        bk = xlrd.open_workbook(path, ignore_workbook_corruption=True)
        parts = []
        for name in bk.sheet_names():
            sh = bk.sheet_by_name(name)
            parts.append("[SHEET %s] rows=%d" % (name, sh.nrows))
            for r in range(sh.nrows):
                vals = []
                for c in range(sh.ncols):
                    v = sh.cell_value(r, c)
                    # math.isfinite 不能省：NaN / inf 进 int() 抛 ValueError，会被下面的
                    # 兜底 except 变成 "!!EXTRACT-FAIL!!" + traceback 写进 text/——
                    # 而那串堆栈有字符数，inventory 会显示「正文 N 字符」像是提取成功了。
                    # 空单元格在 xls 里常以 NaN 出现，是很容易撞上的形态。
                    if isinstance(v, float) and math.isfinite(v) and v == int(v):
                        v = int(v)
                    vals.append(str(v).replace("\n", "\\n").strip())
                parts.append(" | ".join(vals))
        return "\n".join(parts)
    except Exception:
        return "!!EXTRACT-FAIL!!\n" + traceback.format_exc()


def extract_xlsx(path):
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append("[SHEET %s] dims=%s" % (ws.title, ws.dimensions))
        for row in ws.iter_rows(values_only=True):
            cells = ["" if v is None else str(v).replace("\n", " / ").strip() for v in row]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def main():
    if len(sys.argv) < 3:
        sys.exit(__doc__)
    src, work = sys.argv[1], sys.argv[2]
    text_dir = os.path.join(work, "text")
    img_dir = os.path.join(work, "images")
    for d in (work, text_dir, img_dir):
        os.makedirs(d, exist_ok=True)

    rows = []
    all_links = {}          # url -> [出处]
    used_tags = {}          # tag -> 首个占用它的 rel（产物名唯一性兜底）
    collisions = []
    for root, _, files in os.walk(src):
        for fn in sorted(files):
            fp = os.path.join(root, fn)
            rel = os.path.relpath(fp, src)
            ext = os.path.splitext(fn)[1].lower()
            size = os.path.getsize(fp)
            # tag 是 text/ 与 images/ 下产物的文件名前缀，**必须逐文件唯一**。
            # 曾经只取 basename 的 stem：不同子目录的同名文件（需求A/导出.docx 与
            # 需求B/导出.docx）、以及同名不同扩展名（报告.docx 与 报告.pdf）会算出
            # 同一个 tag，后写的静默覆盖先写的——而 inventory 仍逐行列出全部文件，
            # 账面说"全量无损"、磁盘上已经丢了。本工具用于**删源前**归档，丢一次
            # 不可恢复（2026-08-16 实测：3 个文件只落 1 份，脚本仍报 ✅ 3 文件）。
            # 现在：相对路径入 tag（safe_name 把路径分隔符转成 _）+ 碰撞追加序号并
            # 告警——safe_name 有 80 字截断，深目录长文件名仍可能撞，兜底不能省。
            base_tag = safe_name(os.path.splitext(rel)[0])
            tag = base_tag
            n = 2
            while tag in used_tags:
                tag = "%s-%d" % (base_tag, n)
                n += 1
            if tag != base_tag:
                collisions.append((rel, used_tags[base_tag], tag))
            used_tags[tag] = rel
            body, n_img, note = "", 0, ""
            try:
                if ext == ".docx":
                    body, n_img, links, img_fail = extract_docx(fp, img_dir, tag)
                    if img_fail:
                        note = "%d 张内嵌图写入失败（路径过长/权限），正文已提取" % img_fail
                    for u in links:
                        all_links.setdefault(u, []).append(rel + " (超链接)")
                elif ext == ".pdf":
                    body, n_img = extract_pdf(fp, img_dir, tag)
                elif ext == ".xls":
                    body = extract_xls(fp)
                elif ext in (".xlsx", ".xlsm"):
                    body = extract_xlsx(fp)
                elif ext in (".md", ".txt", ".csv"):
                    body = io.open(fp, encoding="utf-8", errors="replace").read()
                else:
                    note = "二进制/不可转写——归档处置需人工判定"
            except ImportError as e:
                mod = getattr(e, "name", None) or str(e)
                note = "缺依赖库 %s（安装：pip install %s）后重跑" % (mod, PIP_NAMES.get(mod, mod))
            except Exception as e:
                note = "提取失败: %s" % type(e).__name__
            if body:
                with io.open(os.path.join(text_dir, tag + ".txt"), "w",
                             encoding="utf-8", newline="\n") as f:
                    f.write(body)
                for u in URL_RE.findall(body):
                    all_links.setdefault(u.rstrip("。，.;"), []).append(rel + " (正文)")
            rows.append((rel, ext or "?", size, len(body), n_img, note, tag if body else ""))

    # inventory.md
    # 「产物」列是账实对账的锚点：源文件 → 实际落盘的 text/ 文件一一对应。
    # 没有这一列时，产物互相覆盖后台账照样逐行列出全部源文件，看不出磁盘上少了东西。
    lines = ["# 收料模态盘点", "",
             "| 文件 | 类型 | 体量 | 正文字符 | 内嵌图 | 产物 | 备注 |",
             "|---|---|---|---|---|---|---|"]
    for rel, ext, size, n_txt, n_img, note, tag in rows:
        lines.append("| %s | %s | %.0fKB | %d | %d | %s | %s |"
                     % (rel, ext, size / 1024, n_txt, n_img,
                        ("text/%s.txt" % tag) if tag else "—", note))
    lines += ["", "合计 %d 个文件，正文 %d 字符，内嵌图 %d 张，外链 %d 条"
              % (len(rows), sum(r[3] for r in rows), sum(r[4] for r in rows), len(all_links)),
              "", "> 下一步：①逐文件通读 text/ ②逐张评估 images/（100% 留痕）③links.md 逐条判覆盖性"]
    with io.open(os.path.join(work, "inventory.md"), "w", encoding="utf-8", newline="\n") as f:
        f.write("\n".join(lines))

    # links.md
    llines = ["# 外链清单（删源前知情项）", "",
              "| 外链 | 出处 | 覆盖判定(待填) |", "|---|---|---|"]
    for u, srcs in sorted(all_links.items()):
        llines.append("| %s | %s | 待判定 |" % (u, "；".join(sorted(set(srcs))[:3])))
    with io.open(os.path.join(work, "links.md"), "w", encoding="utf-8", newline="\n") as f:
        f.write("\n".join(llines))

    n_out = sum(1 for r in rows if r[6])
    print("✅ 盘点完成：%d 文件 / 正文产物 %d 份 / 图 %d 张 / 外链 %d 条" %
          (len(rows), n_out, sum(r[4] for r in rows), len(all_links)))
    print("   inventory.md / text/ / images/ / links.md → %s" % work)
    # 碰撞必须让人看见：这是「删源前」归档，静默改名比静默覆盖轻，但仍需知情
    if collisions:
        print("⚠ %d 组文件的产物名发生碰撞（同名不同目录 / 同名不同扩展名），已自动加序号区分：" % len(collisions))
        for rel, first, tag in collisions[:5]:
            print("   %s 与 %s 撞名 → 本份产物改用 %s" % (rel, first, tag))
        if len(collisions) > 5:
            print("   …… 另 %d 组，逐条见 inventory.md 的「产物」列" % (len(collisions) - 5))


if __name__ == "__main__":
    main()
